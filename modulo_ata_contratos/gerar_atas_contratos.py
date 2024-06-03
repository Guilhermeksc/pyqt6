from PyQt6.QtWidgets import *
from PyQt6.QtCore import *
from PyQt6.QtGui import *
from pathlib import Path
from modulo_ata_contratos.regex_termo_homolog import *
from modulo_ata_contratos.regex_sicaf import *
from modulo_ata_contratos.processar_homologacao import ProgressDialog
from modulo_ata_contratos.processar_sicaf import SICAFDialog
from modulo_ata_contratos.relatorio_indicadores import RelatorioIndicadores
from modulo_ata_contratos.utils import create_button, load_icons, apply_standard_style, limpar_quebras_de_linha
from modulo_ata_contratos.data_utils import DatabaseDialog, PDFProcessingThread, atualizar_modelo_com_dados, save_to_dataframe, load_file_path, obter_arquivos_txt, ler_arquivos_txt
from modulo_ata_contratos.canvas_gerar_atas import *
from diretorios import *
import geopandas as gpd
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import contextily as ctx
import traceback
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm 

import seaborn as sns
from planejamento.utilidades_planejamento import DatabaseManager

NUMERO_ATA_GLOBAL = None
GERADOR_NUMERO_ATA = None

class CustomTreeView(QTreeView):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.cnpj = ""

    def contextMenuEvent(self, event):
        index = self.indexAt(event.pos())
        if not index.isValid():
            return

        item = self.model().itemFromIndex(index)
        if item and " - " in item.text():
            # Usar QTextDocument para extrair texto puro a partir do HTML
            doc = QTextDocument()
            doc.setHtml(item.text())
            plain_text = doc.toPlainText()

            # Extração do CNPJ sem HTML
            if " - " in plain_text:
                self.cnpj = plain_text.split(' - ')[0]

            menu = QMenu(self)
            copyAction = QAction(f"Copiar CNPJ {self.cnpj}", self)
            copyAction.triggered.connect(lambda: self.copy_cnpj(plain_text))
            menu.addAction(copyAction)
            menu.exec(event.globalPos())

    def copy_cnpj(self, text):
        if " - " in text:
            self.cnpj = text.split(' - ')[0]  # Extrai o CNPJ
        QApplication.clipboard().setText(self.cnpj)
        QToolTip.showText(QCursor.pos(), f"CNPJ {self.cnpj} copiado para área de transferência", self)

    def mousePressEvent(self, event):
        index = self.indexAt(event.pos())
        if index.isValid() and event.button() == Qt.MouseButton.LeftButton:
            # Expande ou colapsa o item clicado
            self.setExpanded(index, not self.isExpanded(index))
            
            # Se o item foi expandido, expanda também o primeiro nível de subitens
            if self.isExpanded(index):
                model = self.model()
                numRows = model.rowCount(index)
                for row in range(numRows):
                    childIndex = model.index(row, 0, index)
                    self.setExpanded(childIndex, True)
                    # Colapsa todos os subníveis abaixo do primeiro nível
                    self.collapseAllChildren(childIndex)

        super().mousePressEvent(event)

    def collapseAllChildren(self, parentIndex):
        """Recursivamente colapsa todos os subníveis de um dado índice."""
        model = self.model()
        numRows = model.rowCount(parentIndex)
        for row in range(numRows):
            childIndex = model.index(row, 0, parentIndex)
            self.collapseAllChildren(childIndex)
            self.setExpanded(childIndex, False)

class HTMLDelegate(QStyledItemDelegate):
    def paint(self, painter, option, index):
        painter.save()
        options = option
        self.initStyleOption(options, index)
        style = options.widget.style() if options.widget else QApplication.style()
        doc = QTextDocument()
        doc.setHtml(options.text)
        options.text = ""
        style.drawControl(QStyle.ControlElement.CE_ItemViewItem, options, painter)
        ctx = QAbstractTextDocumentLayout.PaintContext()

        textRect = style.subElementRect(QStyle.SubElement.SE_ItemViewItemText, options, None)
        painter.translate(textRect.topLeft())
        painter.setClipRect(textRect.translated(-textRect.topLeft()))
        doc.documentLayout().draw(painter, ctx)
        painter.restore()

    def sizeHint(self, option, index):
        options = option
        self.initStyleOption(options, index)
        doc = QTextDocument()
        doc.setHtml(options.text)
        doc.setTextWidth(options.rect.width())
        return QSize(int(doc.idealWidth()), int(doc.size().height()))
    
class GerarAtasWidget(QWidget):
    def __init__(self, icons_dir, parent=None):
        super().__init__(parent)
        self.icons_dir = Path(icons_dir)
        self.buttons = {}
        self.tr_variavel_df_carregado = None 
        self.pdf_dir = Path(PDF_DIR)
        self.txt_dir = Path(TXT_DIR) 
        self.sicaf_dir = Path(SICAF_DIR)
        self.sicaf_txt_dir = Path(SICAF_TXT_DIR)
        self.mapeamento_colunas = self.obter_mapeamento_colunas()
        self.current_dataframe = None
        self.pe_pattern = None
        self.setup_ui()
        self.progressDialog = ProgressDialog(self.pdf_dir, self)
        self.setup_pdf_processing_thread()        
        
    def obter_mapeamento_colunas(self):
        return {
            "Grupo": "grupo",
            "Item": "item_num",
            "Catálogo": "catalogo",
            "Descrição": "descricao_tr",
            "Descrição Detalhada": "descricao_detalhada",
            "Unidade": "unidade",
            "Quantidade": "quantidade",
            "Valor Estimado": "valor_estimado",
            "Valor Homologado": "valor_homologado_item_unitario",
            "Desconto (%)": "percentual_desconto",
            "Valor Estimado Total": "valor_estimado_total_do_item",
            "Valor Homologado Total": "valor_homologado_total_item",
            "Marca Fabricante": "marca_fabricante",
            "Modelo Versão": "modelo_versao",
            "UASG": "uasg",
            "Órgão Responsável": "orgao_responsavel",
            "Número": "num_pregao",
            "Ano": "ano_pregao",
            "SRP": "srp",
            "Objeto": "objeto",
            "Situação": "situacao",
            "Melhor Lance": "melhor_lance",
            "Valor Negociado": "valor_negociado",
            "Ordenador Despesa": "ordenador_despesa",
            "Empresa": "empresa",
            "CNPJ": "cnpj",
            "Endereço": "endereco",
            "CEP": "cep",
            "Município": "municipio",
            "Telefone": "telefone",
            "Email": "email",
            "Responsável Legal": "responsavel_legal"
        }
    
    def setup_ui(self):
        self.main_layout = QVBoxLayout(self)
        self.setup_alert_label()
        self.setup_buttons()
        self.setup_treeview()
        self.setup_buttons_down()
        self.setLayout(self.main_layout)
        self.setMinimumSize(1200, 600)

    def setup_alert_label(self):
        icon_path = str(self.icons_dir / 'alert.png')
        text = (f"<img src='{icon_path}' style='vertical-align: middle;' width='24' height='24'> "
                "Pressione '<b><u>Termo de Referência</u></b>' para adicionar os dados 'Catálogo', "
                "'Descrição' e 'Descrição Detalhada' do Termo de Referência. "
                f"<img src='{icon_path}' style='vertical-align: middle;' width='24' height='24'>")
        self.alert_label = QLabel(text)
        self.alert_label.setStyleSheet("color: white; font-size: 14pt; padding: 5px;")
        self.alert_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.main_layout.addWidget(self.alert_label)
        self.atasDialog = None

    def setup_buttons(self):
        self.buttons_layout = QHBoxLayout()
        self.icons = load_icons(self.icons_dir)
        button_definitions = self.obter_definicoes_botoes()
        for name, icon_key, callback, tooltip, animate in button_definitions:
            icon = self.icons.get(icon_key, None)
            button = create_button(name, icon, callback, tooltip, QSize(40, 40), None)
            self.buttons[name] = button
            self.buttons_layout.addWidget(button)
        self.main_layout.addLayout(self.buttons_layout)

    def obter_definicoes_botoes(self):
        return [
            ("Termo de Referência", 'stats', self.import_tr, "Importe um arquivo .xlsx com 4 colunas com índice 'item_num', 'catalogo', 'descricao_tr' e 'descricao_detalada'.", True),
            ("Termo de Homologação", 'data-collection', self.processar_homologacao, "Faça o download dos termos de homologação e mova para a pasta de processamento dos Termos de Homologação", False),
            ("SICAF", 'sicaf', self.processar_sicaf, "Faça o download do SICAF (Nível I - Credenciamento) e mova para a pasta de processamento do SICAF", False),
            ("Ata / Contrato", 'verify_menu', self.abrir_dialog_atas, "Com o database concluíodo é possível gerar as atas ou contratos", False),
        ]

    def setup_buttons_down(self):
        self.buttons_layout = QHBoxLayout()
        self.icons = load_icons(self.icons_dir)
        button_definitions = self.obter_definicoes_botoes_embaixo()
        for name, icon_key, callback, tooltip, animate in button_definitions:
            icon = self.icons.get(icon_key, None)
            button = create_button(name, icon, callback, tooltip, QSize(40, 40), None)
            self.buttons[name] = button
            self.buttons_layout.addWidget(button)
        self.main_layout.addLayout(self.buttons_layout)

    def obter_definicoes_botoes_embaixo(self):
        return [
            ("Database", 'data-processing', self.update_database, "Salva ou Carrega os dados do Database", False),
            ("Salvar Tabela", 'excel', self.salvar_tabela, "Importe um arquivo .xlsx com 4 colunas com índice 'item_num', 'catalogo', 'descricao_tr' e 'descricao_detalada'.", True),
            ("Indicadores", 'performance', self.indicadores_normceim, "Visualize os indicadores do relatório", False),
            ("Configurações", 'gear_menu', self.processar_sicaf, "Faça o download do SICAF (Nível I - Credenciamento) e mova para a pasta de processamento do SICAF", False),
        ]
    
    def setup_treeview(self):
        self.model = QStandardItemModel()  # Inicializando o modelo
        self.treeView = CustomTreeView()
        self.treeView.setModel(self.model)
        self.main_layout.addWidget(self.treeView)
        self.treeView.header().setDefaultAlignment(Qt.AlignmentFlag.AlignCenter)
        self.treeView.setAnimated(True)  # Facilita a visualização da expansão/colapso
        self.treeView.setUniformRowHeights(True)  # Uniformiza a altura das linhas      
        self.treeView.setItemsExpandable(True)  # Garantir que o botão para expandir esteja visível
        self.treeView.setExpandsOnDoubleClick(False)  # Evita a expansão por duplo clique
        self.setup_treeview_styles()
        
    def setup_treeview_styles(self):
        header = self.treeView.header()
        header.setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
        self.treeView.setStyleSheet("""
            QTreeView {
                background-color: #f9f9f9;
                alternate-background-color: #e0e0e0;
                color: #333;
                font-size: 16px;
                border: 1px solid #ccc;
            }
            QTreeView::item:selected {
                background-color: #b0c4de;
                color: white;
            }
            QTreeView::item:hover {
                background-color: #d3d3d3;
                color: black;
            }
            QHeaderView::section {
                background-color: #d3d3d3;
                padding: 5px;
                border: 1px solid #ccc;
                font-size: 16px;
            }
        """)

    def setup_pdf_processing_thread(self):
        self.processing_thread = PDFProcessingThread(self.pdf_dir, self.txt_dir)
        self.processing_thread.progress_updated.connect(self.progressDialog.update_progress)
        self.processing_thread.processing_complete.connect(self.progressDialog.on_conversion_finished)

    def import_tr(self):
        arquivo, _ = QFileDialog.getOpenFileName(self, "Selecionar arquivo", "", "Excel files (*.xlsx *.xls)")
        if arquivo:
            self.tr_variavel_df_carregado = pd.read_excel(arquivo)
            colunas_relevantes = ["item_num", "catalogo", "descricao_tr", "descricao_detalhada"]
            df_relevante = self.tr_variavel_df_carregado[colunas_relevantes]
            QMessageBox.information(self, "Arquivo Carregado", f"O arquivo '{QFileInfo(arquivo).fileName()}' foi carregado com sucesso!")
            self.atualizar_modelo_com_dados(df_relevante)
            self.atualizar_alerta_apos_importar_tr()

    def atualizar_alerta_apos_importar_tr(self):
        icon_path = str(self.icons_dir / 'confirm.png')
        new_text = (f"<img src='{icon_path}' style='vertical-align: middle;' width='24' height='24'> "
                    "Salve os Termos de Homologação na pasta correta e pressione '<b><u>Termo de Homologação</u></b>' para processar os dados. "
                    f"<img src='{icon_path}' style='vertical-align: middle;' width='24' height='24'>")
        self.alert_label.setText(new_text)

    def atualizar_alerta_apos_processar_homologacao(self):
        icon_path = str(self.icons_dir / 'sicaf.png')
        new_text = (f"<img src='{icon_path}' style='vertical-align: middle;' width='24' height='24'> "
                    "Clique com o botão direito no TreeView para copiar o CNPJ para facilitar a busca do SICAF Nível I. "
                    f"<img src='{icon_path}' style='vertical-align: middle;' width='24' height='24'>")
        self.alert_label.setText(new_text)

    def atualizar_modelo_com_dados(self, df_relevante):
        limpar_quebras_de_linha(df_relevante)
        self.model.clear()
        self.model.setHorizontalHeaderLabels(['Item', 'Catálogo', 'Descrição', 'Descrição Detalhada'])
        for _, row in df_relevante.iterrows():
            item_num = QStandardItem(str(row['item_num']))
            catalogo = QStandardItem(str(row['catalogo']))
            descricao_tr = QStandardItem(str(row['descricao_tr']))
            descricao_detalhada = QStandardItem(str(row['descricao_detalhada']))
            item_num.setEditable(False)
            catalogo.setEditable(False)
            descricao_tr.setEditable(False)
            descricao_detalhada.setEditable(False)
            self.model.appendRow([item_num, catalogo, descricao_tr, descricao_detalhada])
        self.treeView.expandAll()
        for column in range(self.model.columnCount()):
            self.treeView.resizeColumnToContents(column)

    def processar_homologacao(self):
        if not self.pdf_dir.exists():
            QMessageBox.warning(self, "Erro", "Pasta de PDFs não encontrada.")
            return
        pdf_files = list(self.pdf_dir.glob("*.pdf"))
        new_files = [file for file in pdf_files if file not in self.progressDialog.processed_files]
        if not new_files:
            QMessageBox.information(self, "Informação", "Nenhum novo arquivo PDF para processar.")
            return
        total_files = len(new_files)
        self.progressDialog = ProgressDialog(total_files, self.pdf_dir, self)
        self.progressDialog.processing_complete.connect(lambda extracted_data: self.finalizar_processamento_homologacao(extracted_data))
        self.progressDialog.show()

    def finalizar_processamento_homologacao(self, extracted_data):
        self.homologacao_dataframe = save_to_dataframe(extracted_data, self.tr_variavel_df_carregado, DATABASE_DIR, self.current_dataframe)
        
        if self.homologacao_dataframe is not None:
            self.current_dataframe = self.homologacao_dataframe  # Atualiza o DataFrame corrente
            self.update_treeview_with_dataframe(self.homologacao_dataframe)
            self.atualizar_alerta_apos_processar_homologacao()
            return self.current_dataframe  # Retorna o DataFrame atualizado
        else:
            QMessageBox.warning(self, "Erro", "Falha ao salvar os dados.")
            return None  # Retorna None para indicar que o processo falhou

    def processar_sicaf(self):
        if self.current_dataframe is not None:
            dataframe_to_use = self.current_dataframe
        else:
            QMessageBox.warning(self, "Erro", "Primeiro processe a homologação ou carregue os dados do banco de dados.")
            return

        if not self.sicaf_dir.exists():
            QMessageBox.warning(self, "Erro", "Pasta de PDFs não encontrada.")
            return

        self.progressSicafDialog = SICAFDialog(self.sicaf_dir, dataframe_to_use, self)
        # Conecta o sinal a ambos os métodos
        self.progressSicafDialog.processing_complete.connect(self.finalizar_processamento_sicaf)
        self.progressSicafDialog.processing_complete.connect(self.receber_df_final)
        self.progressSicafDialog.show()

    def receber_df_final(self, df_final):
        if isinstance(df_final, pd.DataFrame):
            self.current_dataframe = df_final  # Atualize o DataFrame atual
            print("DataFrame final recebido do SICAF:")
            print(df_final)
            return self.current_dataframe
        else:
            QMessageBox.warning(self, "Erro", "Dados recebidos não são válidos.")

    def finalizar_processamento_sicaf(self, extracted_data):
        if isinstance(extracted_data, pd.DataFrame):
            print("DataFrame resultante do SICAF:")
            print(extracted_data)
            self.update_treeview_with_dataframe(extracted_data)
        else:
            print("Erro: Dados recebidos não são um DataFrame.")
            QMessageBox.warning(self, "Erro", "Os dados recebidos não são válidos.")

    def handle_loaded_data(self, loaded_dataframe, pe_pattern=None):
        if isinstance(loaded_dataframe, pd.DataFrame) and not loaded_dataframe.empty:
            self.current_dataframe = loaded_dataframe  # Atualiza o DataFrame corrente
            self.pe_pattern = pe_pattern  # Armazena o padrão PE identificado
            print(f"DataFrame atualizado e carregado:\n{self.current_dataframe.head()}")
            print(f"Padrão PE identificado: {self.pe_pattern}")
            self.update_treeview_with_dataframe(self.current_dataframe)
        else:
            QMessageBox.warning(self, "Aviso", "Os dados carregados não são um DataFrame válido ou estão vazios.")

    def update_database(self):
        # Sempre abre o diálogo, independentemente da existência de um DataFrame atual
        dialog = DatabaseDialog(self, self.current_dataframe, self.handle_loaded_data)
        dialog.exec()

    def update_progress(self, value):
        if self.progressDialog.isVisible():
            self.progressDialog.progressBar.setValue(value)
        else:
            # Caso a barra de progresso não esteja visível, você pode optar por mostrá-la aqui
            self.progressDialog.show()
            self.progressDialog.progressBar.setValue(value)

    def update_treeview_with_dataframe(self, dataframe):
        if dataframe is None:
            QMessageBox.critical(self, "Erro", "O DataFrame não está disponível para atualizar a visualização.")
            return
        creator = ModeloTreeview(self.icons_dir)
        self.model = creator.criar_modelo(dataframe)
        self.treeView.setModel(self.model)
        self.treeView.setItemDelegate(HTMLDelegate())
        self.treeView.reset()
                    
    def abrir_dialog_atas(self):
        if self.current_dataframe is not None:
            dataframe_to_use = self.current_dataframe
            # Verifica se as colunas desejadas estão presentes no dataframe
            if all(col in dataframe_to_use.columns for col in ['empresa', 'num_pregao', 'ano_pregao']):
                print("Colunas de 'empresa', 'num_pregao', 'ano_pregao' do DataFrame:")
                print(dataframe_to_use[['empresa', 'num_pregao', 'ano_pregao']])
            else:
                print("Alguma das colunas 'empresa', 'num_pregao', 'ano_pregao' não está presente no DataFrame.")
        else:
            dataframe_to_use = None  # Define um valor padrão se não houver dataframe atual
            print("Nenhum DataFrame atual disponível.")

        if self.atasDialog is None or not self.atasDialog.isVisible():
            # Passa corretamente o dataframe como argumento para AtasDialog
            self.atasDialog = AtasDialog(self, pe_pattern=self.pe_pattern, dataframe=dataframe_to_use)
            self.atasDialog.show()
        else:
            self.atasDialog.raise_()
            self.atasDialog.activateWindow()

    def salvar_tabela(self):
        if self.current_dataframe is not None:
            # Define o caminho do arquivo a ser salvo
            arquivo_excel = str(self.pdf_dir / 'TabelaAtual.xlsx')
            # Salva o DataFrame no arquivo Excel
            self.current_dataframe.to_excel(arquivo_excel, index=False)
            # Abre o arquivo Excel
            os.startfile(arquivo_excel)
        else:
            QMessageBox.warning(self, "Aviso", "Não há dados para salvar.")

    def indicadores_normceim(self):
        if self.current_dataframe is not None:
            # Supondo que pe_pattern é armazenado em algum lugar após ser determinado
            self.dialogo_indicadores = RelatorioIndicadores(dataframe=self.current_dataframe, parent=self, pe_pattern=self.pe_pattern)
            self.dialogo_indicadores.show()
        else:
            QMessageBox.warning(self, "Aviso", "Não há dados carregados.")

class ModeloTreeview:
    def __init__(self, icons_dir):
        self.check_icon = QIcon(str(icons_dir / 'checked.png'))
        self.uncheck_icon = QIcon(str(icons_dir / 'unchecked.png'))
        self.alert_icon = QIcon(str(icons_dir / 'alert.png'))

    def criar_modelo(self, dataframe):
        model, header = self.initializar_modelo(dataframe)
        empresa_items = self.processar_linhas(dataframe, model)
        self.atualizar_contador_cabecalho(empresa_items, model)
        return model

    def initializar_modelo(self, dataframe):
        total_items = len(dataframe)
        situacoes_analizadas = ['Adjudicado e Homologado', 'Fracassado e Homologado', 'Deserto e Homologado', 'Anulado e Homologado']
        nao_analisados = len(dataframe[~dataframe['situacao'].isin(situacoes_analizadas)])
        header = f"Total de itens da licitação {total_items} | Total de itens analisados {total_items - nao_analisados} | Total de itens não analisados {nao_analisados}"
        model = QStandardItemModel()
        model.setHorizontalHeaderLabels([header])
        return model, header

    def processar_linhas(self, dataframe, model):
        empresa_items = {}
        for _, row in dataframe.iterrows():
            self.processar_linhas_individualmente(row, model, empresa_items)
        return empresa_items

    def processar_linhas_individualmente(self, row, model, empresa_items):
        parent_key, parent_item = self.determinar_itens_iguais(row, empresa_items)
        if parent_item:
            model.appendRow(parent_item)
            if parent_key not in empresa_items:
                empresa_items[parent_key] = {
                    'item': parent_item,
                    'count': 0,
                    'details_added': False,
                    'items_container': QStandardItem()  # Não defina o texto aqui
                }
                empresa_items[parent_key]['items_container'].setEditable(False)

        # Ajuste para incrementar a contagem em todas as situações
        if parent_key in empresa_items:
            empresa_items[parent_key]['count'] += 1

        self.adicionar_informacao_ao_item(row, empresa_items[parent_key]['item'], empresa_items, parent_key)

    def adicionar_informacao_ao_item(self, row, parent_item, empresa_items, parent_key):
        font_size = "14px"  # Define o tamanho da fonte
        situacoes_especificas = ['Fracassado e Homologado', 'Deserto e Homologado', 'Anulado e Homologado']
        situacao = row['situacao']

        # Determinar se a situação é específica ou "Não definido"
        if situacao not in situacoes_especificas and situacao != 'Adjudicado e Homologado':
            situacao = 'Não definido'

        if situacao in situacoes_especificas or situacao == 'Não definido':
            # Cria um item com informações básicas sem detalhes extras para situações específicas
            item_text = f"<span style='font-size: {font_size};'>Item {row['item_num']} - {row['descricao_tr']} - {situacao}</span>"
            item_info = QStandardItem(item_text)
            item_info.setEditable(False)
            parent_item.appendRow(item_info)
        else:
            # Processo normal para 'Adjudicado e Homologado'
            if not empresa_items[parent_key]['details_added']:
                self.adicionar_detalhes_empresa(row, parent_item)
                empresa_items[parent_key]['items_container'].setText("")  # Limpa o texto se necessário
                parent_item.appendRow(empresa_items[parent_key]['items_container'])
                empresa_items[parent_key]['details_added'] = True

            # Adicionando itens específicos da licitação
            self.adicionar_subitens_detalhados(row, empresa_items[parent_key]['items_container'])

        # Atualizar o texto do container com base na contagem de itens
        item_count_text = "Item" if empresa_items[parent_key]['count'] == 1 else "Relação de itens:"
        empresa_items[parent_key]['items_container'].setText(f"<span style='font-size: {font_size};'><b>{item_count_text}</b> ({empresa_items[parent_key]['count']})</span>")
        
    def atualizar_contador_cabecalho(self, empresa_items, model):
        font_size = "16px"  # Definir o tamanho da fonte para os cabeçalhos dos itens
        for chave_item_pai, empresa in empresa_items.items():
            count = empresa['count']
            # Formatar o texto com HTML para ajustar o tamanho da fonte
            display_text = f"<span style='font-size: {font_size};'>{chave_item_pai} (<b>1 item</b>)</span>" if count == 1 else f"<span style='font-size: {font_size};'>{chave_item_pai} (<b>{count} itens</b>)</span>"
            empresa['item'].setText(display_text)

    def adicionar_detalhes_empresa(self, row, parent_item):
        font_size = "14px"
        infos = [
            f"<span style='font-size: {font_size};'><b>Endereço:</b> {row['endereco']}, CEP: {row['cep']}, Município: {row['municipio']}</span>" if pd.notna(row['endereco']) else f"<span style='font-size: {font_size};'><b>Endereço:</b> Não informado</span>",
            f"<span style='font-size: {font_size};'><b>Contato:</b> {row['telefone']} <b>Email:</b> {row['email']}</span>" if pd.notna(row['telefone']) else f"<span style='font-size: {font_size};'><b>Contato:</b> Não informado</span>",
            f"<span style='font-size: {font_size};'><b>Responsável Legal:</b> {row['responsavel_legal']}</span>" if pd.notna(row['responsavel_legal']) else f"<span style='font-size: {font_size};'><b>Responsável Legal:</b> Não informado</span>"
        ]
        for info in infos:
            info_item = QStandardItem(info)
            info_item.setEditable(False)
            parent_item.appendRow(info_item)

    def determinar_itens_iguais(self, row, empresa_items):
        empresa_name = str(row['empresa']) if pd.notna(row['empresa']) else ""
        cnpj = str(row['cnpj']) if pd.notna(row['cnpj']) else ""
        situacao = str(row['situacao']) if pd.notna(row['situacao']) else "Não definido"
        is_situacao_only = not empresa_name and not cnpj
        parent_key = f"{situacao}" if is_situacao_only else f"{cnpj} - {empresa_name}".strip(" -")
        if parent_key not in empresa_items:
            parent_item = QStandardItem()
            parent_item.setEditable(False)
            parent_item.setIcon(self.alert_icon if is_situacao_only else (self.check_icon if pd.notna(row['endereco']) else self.uncheck_icon))
            return parent_key, parent_item
        return parent_key, None

    def criar_dados_sicaf_do_item(self, row):
        fields = ['endereco', 'cep', 'municipio', 'telefone', 'email', 'responsavel_legal']
        return [self.criar_detalhe_item(field.capitalize(), row[field]) for field in fields if pd.notna(row[field])]

    def adicionar_subitens_detalhados(self, row, sub_items_layout):
        font_size = "14px"  # Definir o tamanho da fonte para os detalhes dos itens
        # Criar o item principal com formatação HTML para o tamanho da fonte
        item_info_html = f"<span style='font-size: {font_size};'>Item {row['item_num']} - {row['descricao_tr']} - {row['situacao']}</span>"
        item_info = QStandardItem(item_info_html)
        item_info.setEditable(False)
        sub_items_layout.appendRow(item_info)

        # Adicionar mais detalhes com formatação HTML
        detalhes_html = [
            f"<span style='font-size: {font_size};'><b>Descrição Detalhada:</b> {row['descricao_detalhada']}</span>",
            f"<span style='font-size: {font_size};'><b>Unidade de Fornecimento:</b> {row['unidade']} <b>Quantidade:</b> {self.formatar_quantidade(row['quantidade'])} <b>Valor Estimado:</b> {self.formatar_brl(row['valor_estimado'])} <b>Valor Homologado:</b> {self.formatar_brl(row['valor_homologado_item_unitario'])} <b>Desconto:</b> {self.formatar_percentual(row['percentual_desconto'])} <b>Marca:</b> {row['marca_fabricante']} <b>Modelo:</b> {row['modelo_versao']}</span>",
        ]

        for detalhe_html in detalhes_html:
            detalhe_item = QStandardItem(detalhe_html)
            detalhe_item.setEditable(False)
            item_info.appendRow(detalhe_item)


    def criar_detalhe_item(self, label, data):
        return QStandardItem(f"<b>{label}:</b> {data if pd.notna(data) else 'Não informado'}")

    def formatar_brl(self, valor):
        try:
            if valor is None:
                return "Não disponível"  # ou outra representação adequada para seu caso de uso
            return f"R$ {float(valor):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        except ValueError:
            return "Valor inválido"

    def formatar_quantidade(self, valor):
        try:
            float_value = float(valor)
            if float_value.is_integer():
                return f"{int(float_value)}"
            else:
                return f"{float_value:.2f}".replace('.', ',')
        except ValueError:
            return "Erro de Formatação"

    def formatar_percentual(self, valor):
        try:
            percent_value = float(valor)
            return f"{percent_value:.2f}%"
        except ValueError:
            return "Erro de Formatação"
        
class AtasDialog(QDialog):
    NUMERO_ATA_GLOBAL = None  # Defina isso em algum lugar adequado dentro de sua classe

    def __init__(self, parent=None, pe_pattern=None, dataframe=None):
        super().__init__(parent)
        self.db_manager = DatabaseManager(CONTROLE_DADOS)
        self.pe_pattern = pe_pattern
        self.nup_data = None
        self.dataframe = dataframe 
        self.setWindowTitle("Geração de Atas / Contratos")
        self.setFont(QFont('Arial', 14))
        layout = QVBoxLayout(self)

        # Primeiro crie a QLabel para o último contrato
        self.ultimo_contrato_label = QLabel("O último contrato gerado foi:")
        self.ultimo_contrato_label.setFont(QFont('Arial', 14))
        layout.addWidget(self.ultimo_contrato_label)

        self.label = QLabel("\nDigite o próximo Número de Controle de Atas/Contratos:\n")
        self.label.setFont(QFont('Arial', 14))
        layout.addWidget(self.label)

        # Cria um QHBoxLayout para a entrada e o botão
        entry_button_layout = QHBoxLayout()

        # Espaçador à esquerda
        left_spacer = QSpacerItem(40, 20, QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)

        entry_button_layout.addItem(left_spacer)

        # Agora crie o QLineEdit para a entrada de texto
        self.ataEntry = QLineEdit(self)
        self.ataEntry.setPlaceholderText("Digite um número até 4 dígitos")
        self.ataEntry.setMaxLength(4)

        self.ataEntry.setFixedWidth(self.ataEntry.fontMetrics().horizontalAdvance('0') * 6)

        entry_button_layout.addWidget(self.ataEntry)

        # Cria o botão Confirmar
        self.confirmButton = QPushButton("Confirmar", self)
        self.confirmButton.clicked.connect(self.confirmar_numero_ata_e_nup_do_processo)
        entry_button_layout.addWidget(self.confirmButton)

        # Espaçador à direita
        right_spacer = QSpacerItem(40, 20, QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)

        entry_button_layout.addItem(right_spacer)
      
        # Adiciona o QHBoxLayout da entrada e botão ao layout principal
        layout.addLayout(entry_button_layout)

        # Cria um QHBoxLayout para os botões Gerar Atas/Contratos e Gerar Documento
        buttons_layout = QHBoxLayout()

        self.label_espaco = QLabel("\n")
        layout.addWidget(self.label_espaco)

        # Botão para gerar atas ou contratos
        self.gerarButton = self.criar_botao_especial("Gerar\nAtas", str(ICONS_DIR / 'gerar_atas.png'))
        self.gerarButton.clicked.connect(self.gerar_atas_contratos)
        buttons_layout.addWidget(self.gerarButton)

        # Botão para gerar documento
        self.gerarDocumentoButton = self.criar_botao_especial("Gerar\nContratos", str(ICONS_DIR / 'gerar_contrato.png'))
        self.gerarDocumentoButton.clicked.connect(self.gerar_documento)
        buttons_layout.addWidget(self.gerarDocumentoButton)
        
        # Depois de criar self.ataEntry, agora você pode verificar e definir seu valor inicial
        ultimo_num_contrato = self.carregar_ultimo_contrato()
        if ultimo_num_contrato is not None:
            self.atualizar_ultimo_contrato_label(f"Nº {ultimo_num_contrato}")
            self.ataEntry.setText(str(ultimo_num_contrato + 1))
        else:
            self.ultimo_contrato_label.setText("O último número de ata/contrato gerado foi: Nenhum")
        
        # Adiciona o QHBoxLayout dos botões ao layout principal
        layout.addLayout(buttons_layout)

    @staticmethod
    def convert_pe_format(pe_string):
        pe_formatted = pe_string.replace('PE-', 'PE ').replace('-', '/')
        print(f"Converted PE format: {pe_formatted}")  # Depuração
        return pe_formatted

    def obter_nup(self, pe_formatted):
        try:
            with self.db_manager as conn:
                query = f"SELECT nup FROM controle_processos WHERE id_processo LIKE '%{pe_formatted}%'"
                df = pd.read_sql(query, conn)
                if not df.empty:
                    self.nup_data = {
                        'nup': df.iloc[0]['nup']
                    }
                    return self.nup_data
                else:
                    return None
        except Exception as e:
            print(f"Erro ao acessar o banco de dados: {e}")
            return None
                    
    def atualizar_ultimo_contrato_label(self, ultimo_num_contrato):
        self.ultimo_contrato_label.setText(f"O último número de ata/contrato gerado foi: {ultimo_num_contrato}")

    def salvar_ultimo_contrato(self, ultimo_num_contrato):
        with open(ULTIMO_CONTRATO_DIR, "w") as f:
            f.write(str(ultimo_num_contrato))  # Convertendo para string

    def carregar_ultimo_contrato(self):
        try:
            with open(ULTIMO_CONTRATO_DIR, "r") as f:
                return int(f.read().split('/')[-1])
        except (FileNotFoundError, ValueError):
            return None

    def confirmar_numero_ata_e_nup_do_processo(self):
        numero_ata = self.ataEntry.text()
        if numero_ata.isdigit() and len(numero_ata) <= 4:
            AtasDialog.NUMERO_ATA_GLOBAL = int(numero_ata)
            self.nup_data = self.obter_nup(self.convert_pe_format(self.pe_pattern))
            QMessageBox.information(self, "Número Confirmado", f"Número da ata definido para: {numero_ata}")
        else:
            QMessageBox.warning(self, "Número Inválido", "Por favor, digite um número válido de até 4 dígitos.")

    def criar_botao_especial(self, text, icon_path):
        button = QToolButton(self)
        button.setText(text)
        button.setIcon(QIcon(icon_path))
        button.setIconSize(QSize(64, 64))  # Defina o tamanho do ícone conforme necessário
        button.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextUnderIcon)

        button.setFixedSize(200, 160) 
        return button

    def gerar_atas_contratos(self):
        if not self.nup_data:  # Verifica se nup_data está vazia ou é None
            self.nup_data = "(INSIRA O NUP)"  # Atribui um valor padrão caso não exista nup_data
        self.iniciar_processo(self.nup_data, self.dataframe)

    def iniciar_processo(self, nup_data, dataframe):
        if AtasDialog.NUMERO_ATA_GLOBAL is None:
            raise ValueError("O número da ATA não foi definido!")

        # Chama as outras funções que dependem de NUMERO_ATA_GLOBAL
        criar_pastas_com_subpastas(dataframe)
        ultimo_num_ata = processar_ata(AtasDialog.NUMERO_ATA_GLOBAL, nup_data, dataframe)

        # Atualizar e salvar o último número da ATA
        self.salvar_ultimo_contrato(ultimo_num_ata)
        self.atualizar_ultimo_contrato_label(ultimo_num_ata)

    def gerar_documento(self):
        # Aqui chamamos a função iniciar_processo
        try:
            self.iniciar_contrato()
        except ValueError as e:
            QMessageBox.critical(self, "Erro", str(e))
        pass

    def iniciar_contrato(self):
        if AtasDialog.NUMERO_ATA_GLOBAL is None:
            raise ValueError("O número do Contrato não foi definido!")

        # Chama as outras funções que dependem de NUMERO_ATA_GLOBAL
        criar_pastas_com_subpastas()
        ultimo_num_contrato = processar_contrato(AtasDialog.NUMERO_ATA_GLOBAL)

        # Atualizar e salvar o último número do contrato
        self.salvar_ultimo_contrato(ultimo_num_contrato)
        self.atualizar_ultimo_contrato_label(ultimo_num_contrato)